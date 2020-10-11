from django.shortcuts import render
from django.http import HttpResponse
import requests

from docx import Document
from docx.shared import Inches
import zipfile
from io import BytesIO
from io import StringIO
from cc_script import Copycase

from requests_html import HTML, HTMLSession
from bs4 import BeautifulSoup
import time
import random
from django.shortcuts import redirect

from apiclient.discovery import build

def main(request):
    return render(request, 'main.html')

def get_cases_g(request):
    # GET FROM SEARCH
    nitems = 0
    api_key = 'AIzaSyBHM12AjMw07SlegaQmCiqVgsdUHwfpJgU'
    if request.method == 'GET':
        gr_in = request.GET.get('sbox')
        siterefs = request.GET.getlist('srchk')
    if 'lawphil.net' in siterefs:
        siterefs.append('www.lawphil.net')
    print(gr_in)
    print(siterefs)
    if gr_in == '':
        print('Empty Input')
        epm_txt = "It's empty........ well at least you tried"
        f = BytesIO()
        zip_arch = zipfile.ZipFile(f, 'a')
        doc = Document()
        doc.add_paragraph(epm_txt)
        docx_title = "Empty Search.docx"
        doc.save(docx_title)
        print('Docx has been saved')
        zip_arch.write(docx_title)
    else:
        # MAIN ALGO
        # gr_in = '164457, 201302'
        f = BytesIO()
        zip_arch = zipfile.ZipFile(f, 'a')
        userque = Copycase(gr_in)
        userque.sort_que()
        userque.rev_list()
        # siterefs = ['lawphil.net', 'www.chanrobles.com']

        # START
        userque.nitems = 0
        uql = userque.uquelist
        for query in uql:
            query = query.upper()
            try:
                resource = build("customsearch", 'v1', developerKey = api_key).cse()
                result = resource.list(q=query, cx='6b76bdba34f7fc3eb').execute()
                clink_list = []
                for item in result["items"]:
                    l_title = item["title"]
                    clinkx = item["link"]
                    if query in l_title:
                        print(f'title: {l_title}')
                        print(f'Link: {clinkx}')
                        clink_list.append(clinkx)
                print(clink_list)
                if clink_list == []:
                    qstate = 0  # QUERY HAS NO MATCH
                else:
                    qstate = 1
                for qlinkx in clink_list:
                    try:
                        print(f'onlink: {qlinkx}')
                        qstate = 1
                        rs = HTMLSession()
                        r = rs.get(qlinkx)
                        doc = Document()
                        rlinkx = qlinkx.split('/')[2]
                        # -----LAWPHIL-----
                        if rlinkx == 'lawphil.net' or rlinkx == 'www.lawphil.net':
                            print('try lawphil')
                            try:
                                parags = r.html.find('p')
                                haveTitle = False
                                titleMatch = False
                                for p in parags:
                                    aline = p.find('p', first=True).text
                                    if not (haveTitle) and 'G.R.' in aline:
                                        if '\n' in aline:  # IF HAVE NEW LINE
                                            # print('NEWLINE')
                                            tt = aline.split('\n')
                                            for x in range(len(tt)):
                                                if query in tt[x]:
                                                    title = tt[x] + ' ' + tt[len(tt) - 1]
                                                    break
                                        else:
                                            # print('TABS/SPACE')
                                            if '\xa0' in aline:  # IF HAVE TAB
                                                tt = aline.split('\xa0')
                                                title = tt[0] + ' ' + tt[len(tt) - 1]
                                            else:
                                                title = aline.upper()
                                        stitle = title.split(' ')
                                        for x in range(len(stitle)):
                                            if query in stitle[x]:
                                                qstate = 1
                                                titleMatch = True
                                                print(f"{title.split(' ')} - match")
                                                break
                                            else:
                                                print(f"{title.split(' ')} - not match")
                                                qstate = 0
                                        haveTitle = True
                                        print(f'Found: {title}')
                                    doc.add_paragraph(aline)
                            except:
                                # Query not found
                                pass
                            if titleMatch:
                                docx_title = 'lawphil - ' + title + '.docx'
                                doc.save(docx_title)
                                print('Docx has been saved')
                                zip_arch.write(docx_title)
                                print('Added to Zip')
                                nitems += 1

                        # -----CHANROBLES-----
                        elif rlinkx == 'www.chanrobles.com':
                            try:
                                print('try chanrobles')
                                parags = r.html.find('.content')
                                th = r.html.find('title')
                                ftit = th[0].text.split('-')
                                print(ftit[0][len(ftit[0]) - 1])
                                print(ftit)
                                if ftit[0][len(ftit[0]) - 1] == 'L':
                                    if ftit[2][:len(ftit[2]) - 2].isnumeric():
                                        title = ftit[0] + '-' + ftit[1] + '-' + ftit[2]
                                    else:
                                        title = ftit[0] + '-' + ftit[1]
                                else:
                                    title = ftit[0]
                                title = title.upper()
                                print(f'found title: {title}')
                                if parags == []:
                                    try:
                                        print('try chanrobles2')
                                        parags = r.html.find('p')
                                        titleMatch = False
                                        for p in parags:
                                            aline = p.find('p', first=True).text
                                            if 'CLICK' in aline.split(' ')[0].upper():
                                                pass
                                            else:
                                                doc.add_paragraph(aline)
                                        stitle = title.split(' ')
                                        for x in range(len(stitle)):
                                            if query in stitle[x]:
                                                qstate = 1
                                                titleMatch = True
                                                print(f"{title.split(' ')} - match")
                                                break
                                            else:
                                                qstate = 0
                                                print(f"{title.split(' ')} - not match")
                                    except:
                                        # Query not found
                                        print('query not found')
                                        pass
                                else:
                                    titleMatch = False
                                    for p in parags:
                                        aline = p.find('p', first=True).text
                                        if aline.split(' ')[0].upper() == 'CLICK':
                                            pass
                                        else:
                                            doc.add_paragraph(aline)
                                    stitle = title.split(' ')
                                    for x in range(len(stitle)):
                                        if query in stitle[x]:
                                            qstate = 1
                                            titleMatch = True
                                            print(f"{title.split(' ')} - match")
                                            break
                                        else:
                                            qstate = 0
                                            print(f"{title.split(' ')} - not match")
                            except:
                                print('query not found')
                                pass
                            if titleMatch:
                                qstate = 1
                                docx_title = 'c-robles - ' + title + '.docx'
                                doc.save(docx_title)
                                print('Docx has been saved')
                                zip_arch.write(docx_title)
                                print('Added to Zip')
                                nitems += 1
                        else:
                            # reference not recog
                            pass
                    except:
                        qstate = 2
                        pass
            except:
                pass
    zip_arch.close()
    zip_title = "CopiedCases.zip"
    response = HttpResponse(
        f.getvalue(),
        content_type='application/zip'
    )
    response['Content-Disposition'] = 'attachment; filename=' + zip_title
    return response
